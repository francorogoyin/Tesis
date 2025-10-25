#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar documento Word con resumen de análisis estadísticos
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def crear_documento():
    """Crea el documento Word con todas las secciones"""

    # Crear documento
    doc = Document()

    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ===========================================================================
    # PORTADA
    # ===========================================================================
    title = doc.add_heading('Resumen de Análisis Estadísticos', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Congruencia Ideológica, Filtrados y Modelos SEM')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_page_break()

    print("Portada creada...")

    # ===========================================================================
    # SECCIÓN 1: CO Y CT CONGRUENTES E INCONGRUENTES
    # ===========================================================================

    doc.add_heading('1. Congruencia e Incongruencia Ideológica', 1)

    # Marco conceptual
    doc.add_heading('1.1. Marco Conceptual', 2)

    p = doc.add_paragraph()
    p.add_run('Definición de congruencia ideológica: ').bold = True
    p.add_run(
        'Se define como la correspondencia entre el tipo de ítem (progresista o conservador) '
        'y la dirección del cambio de opinión o tiempo hacia candidatos ideológicamente afines.'
    )

    # Lista de definiciones
    doc.add_paragraph('Congruente:', style='List Bullet')
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run('Ítems Progresistas + movimiento hacia candidatos de Izquierda')

    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run('Ítems Conservadores + movimiento hacia candidatos de Derecha')

    doc.add_paragraph('Incongruente:', style='List Bullet')
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run('Ítems Progresistas + movimiento hacia candidatos de Derecha')

    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run('Ítems Conservadores + movimiento hacia candidatos de Izquierda')

    # Variables creadas
    doc.add_heading('1.2. Variables Creadas', 2)

    p = doc.add_paragraph()
    p.add_run('Se crearon cuatro variables sumadas para cada elección (Generales y Ballotage):')

    doc.add_paragraph('CO_Congruente = Cambio_Op_Sum_Pro_Izq + Cambio_Op_Sum_Con_Der', style='List Bullet')
    doc.add_paragraph('CO_Incongruente = Cambio_Op_Sum_Pro_Der + Cambio_Op_Sum_Con_Izq', style='List Bullet')
    doc.add_paragraph('CT_Congruente = Cambio_Tiempo_Sum_Pro_Izq + Cambio_Tiempo_Sum_Con_Der', style='List Bullet')
    doc.add_paragraph('CT_Incongruente = Cambio_Tiempo_Sum_Pro_Der + Cambio_Tiempo_Sum_Con_Izq', style='List Bullet')

    # Análisis estadístico
    doc.add_heading('1.3. Análisis Estadístico', 2)

    p = doc.add_paragraph()
    p.add_run('Metodología: ').bold = True
    p.add_run(
        'Se utilizó el test de Wilcoxon pareado para comparar variables congruentes vs incongruentes. '
        'Este test no paramétrico es apropiado para datos pareados que pueden no seguir una distribución normal.'
    )

    # Resultados Generales - Todas las poblaciones
    doc.add_heading('1.4. Resultados: Análisis Global', 2)

    # Tabla con resultados
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Light Grid Accent 1'

    # Headers
    headers = ['Dataset', 'Tipo', 'n', 'Media Congruente', 'Media Incongruente', 'p-valor']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Datos
    data = [
        ['Generales', 'CO', '2786', '0.7968', '-0.7175', '<0.001 ***'],
        ['Generales', 'CT', '2786', '-20.18', '-18.76', '<0.001 ***'],
        ['Ballotage', 'CO', '1254', '0.3373', '-0.2241', '0.009 **'],
        ['Ballotage', 'CT', '1254', '-18.37', '-16.83', '<0.001 ***']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    p = doc.add_paragraph()
    p.add_run('Interpretación: ').bold = True
    p.add_run(
        'Todas las comparaciones resultaron estadísticamente significativas (p < 0.05). '
        'En ambas elecciones, el Cambio de Opinión (CO) fue significativamente mayor en la dirección congruente, '
        'indicando que los participantes tendieron a cambiar sus opiniones en direcciones ideológicamente consistentes. '
        'Para Cambio de Tiempo (CT), también se observaron diferencias significativas, aunque con valores negativos '
        'indicando tiempos de respuesta más rápidos en general.'
    )

    print("Sección 1 creada...")

    # Resultados por categoría
    doc.add_heading('1.5. Resultados por Categoría Ideológica', 2)

    p = doc.add_paragraph()
    p.add_run('Resultados destacados en Generales:')

    doc.add_paragraph('Left Wing: CO (p < 0.001) y CT (p < 0.001) significativos', style='List Bullet')
    doc.add_paragraph('Progressivism: CO (p < 0.001) y CT (p < 0.001) significativos', style='List Bullet')
    doc.add_paragraph('Right Wing Libertarian: CO (p = 0.002) y CT (p = 0.004) significativos', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Resultados destacados en Ballotage:')

    doc.add_paragraph('Progressivism: CO (p = 0.009) y CT (p < 0.001) significativos', style='List Bullet')
    doc.add_paragraph('Moderate Right B: Solo CT significativo (p = 0.034)', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Conclusión: ').bold = True
    p.add_run(
        'El efecto de congruencia ideológica es más robusto en Generales (11 de 12 comparaciones significativas por categoría) '
        'que en Ballotage (3 de 12), sugiriendo que la congruencia ideológica juega un rol más determinante en la '
        'primera vuelta electoral.'
    )

    # Ubicación de archivos
    doc.add_heading('1.6. Archivos Generados', 2)

    doc.add_paragraph('/Data/Procesados/Generales_con_Congruencia.xlsx', style='List Bullet')
    doc.add_paragraph('/Data/Procesados/Ballotage_con_Congruencia.xlsx', style='List Bullet')
    doc.add_paragraph('/Data/Procesados/Resultados_Congruencia_General.xlsx', style='List Bullet')
    doc.add_paragraph('/Data/Procesados/Resultados_Congruencia_Por_Categoria.xlsx', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Notebook fuente: ').bold = True
    p.add_run('52. Análisis Congruencia Ideológica.ipynb')

    doc.add_page_break()

    print("Continuando con Sección 2...")

    # ===========================================================================
    # SECCIÓN 2: FILTRADO CRUZADO
    # ===========================================================================

    doc.add_heading('2. Filtrado Cruzado entre Elecciones', 1)

    # Marco conceptual
    doc.add_heading('2.1. Marco Conceptual', 2)

    p = doc.add_paragraph()
    p.add_run('Objetivo: ').bold = True
    p.add_run(
        'Validar la robustez de los ítems significativos identificados mediante el test de Kruskal-Wallis, '
        'utilizando una estrategia de validación cruzada entre elecciones.'
    )

    p = doc.add_paragraph()
    p.add_run('Metodología: ').bold = True
    p.add_run(
        'Tradicionalmente, cada elección filtra sus datos usando sus propios ítems significativos. '
        'El filtrado cruzado invierte esta lógica:'
    )

    doc.add_paragraph('Ballotage filtrado con ítems significativos de Generales', style='List Bullet')
    doc.add_paragraph('Generales filtrado con ítems significativos de Ballotage', style='List Bullet')

    doc.add_heading('2.2. Ítems Significativos Identificados', 2)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = ['Elección', 'CO Significativos', 'CT Significativos']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ['Generales', '11 ítems', '14 ítems'],
        ['Ballotage', '9 ítems', '5 ítems']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    p = doc.add_paragraph()
    p.add_run('Nota: ').italic = True
    p.add_run(
        'Los ítems significativos fueron determinados mediante test de Kruskal-Wallis con p < 0.05, '
        'comparando las 6 categorías ideológicas válidas.'
    )

    doc.add_heading('2.3. Variables Creadas', 2)

    p = doc.add_paragraph()
    p.add_run('Se crearon 16 nuevas variables de filtrado cruzado:')

    doc.add_paragraph('En Ballotage (usando filtro de Generales):', style='List Bullet')
    doc.add_paragraph('4 variables CO: Pro_Izq, Pro_Der, Con_Izq, Con_Der', style='List Bullet 2')
    doc.add_paragraph('4 variables CT: Pro_Izq, Pro_Der, Con_Izq, Con_Der', style='List Bullet 2')

    doc.add_paragraph('En Generales (usando filtro de Ballotage):', style='List Bullet')
    doc.add_paragraph('4 variables CO: Pro_Izq, Pro_Der, Con_Izq, Con_Der', style='List Bullet 2')
    doc.add_paragraph('4 variables CT: Pro_Izq, Pro_Der, Con_Izq, Con_Der', style='List Bullet 2')

    doc.add_heading('2.4. Resultados: Comparación Filtro Original vs Cruzado', 2)

    p = doc.add_paragraph()
    p.add_run('Análisis estadístico: ').bold = True
    p.add_run('Test de Wilcoxon pareado comparando variables filtradas originales vs cruzadas.')

    p = doc.add_paragraph()
    p.add_run('Hallazgos principales:')

    doc.add_paragraph('14 de 16 comparaciones mostraron diferencias significativas (p < 0.05)', style='List Bullet')
    doc.add_paragraph(
        'Las variables CT mostraron mayor sensibilidad al filtro utilizado que las variables CO',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Las diferencias sugieren que los ítems significativos difieren sustancialmente entre elecciones',
        style='List Bullet'
    )

    p = doc.add_paragraph()
    p.add_run('Interpretación: ').bold = True
    p.add_run(
        'La alta proporción de diferencias significativas (87.5%) entre filtros originales y cruzados '
        'indica que los ítems que discriminan entre categorías ideológicas son específicos de cada '
        'contexto electoral, sugiriendo que diferentes temas adquieren relevancia en Generales vs Ballotage.'
    )

    doc.add_heading('2.5. Archivos Generados', 2)

    doc.add_paragraph('/Data/Procesados/df_Elecciones.xlsx (actualizado)', style='List Bullet')
    doc.add_paragraph('/Data/Procesados/Resultados_Filtrado_Cruzado.xlsx', style='List Bullet')
    doc.add_paragraph('/Data/Procesados/Items_Significativos_Por_Eleccion.xlsx', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Notebook fuente: ').bold = True
    p.add_run('53. Filtrado Cruzado entre Elecciones.ipynb')

    doc.add_page_break()

    print("Continuando con Sección 3...")

    # ===========================================================================
    # SECCIÓN 3: CT FILTRADOS
    # ===========================================================================

    doc.add_heading('3. Cambio de Tiempo Filtrado (CT Filtrado)', 1)

    doc.add_heading('3.1. Marco Conceptual', 2)

    p = doc.add_paragraph()
    p.add_run('Justificación: ').bold = True
    p.add_run(
        'No todos los ítems del cuestionario mostraron diferencias significativas entre categorías ideológicas. '
        'El filtrado permite concentrar el análisis únicamente en aquellos ítems que efectivamente discriminan '
        'entre grupos, aumentando la señal y reduciendo el ruido.'
    )

    p = doc.add_paragraph()
    p.add_run('Criterio de filtrado: ').bold = True
    p.add_run(
        'Test de Kruskal-Wallis con p < 0.05, comparando las 6 categorías ideológicas. '
        'Este test no paramétrico es apropiado dado que los tiempos de respuesta típicamente '
        'no siguen distribución normal.'
    )

    doc.add_heading('3.2. Proceso de Filtrado', 2)

    doc.add_paragraph(
        'Paso 1: Ejecutar Kruskal-Wallis para cada uno de los 40 ítems CT (20 ítems × 2 candidatos)',
        style='List Number'
    )
    doc.add_paragraph(
        'Paso 2: Identificar ítems con p < 0.05',
        style='List Number'
    )
    doc.add_paragraph(
        'Paso 3: Crear variables sumadas solo con ítems significativos por categoría cruzada',
        style='List Number'
    )
    doc.add_paragraph(
        'Paso 4: Analizar diferencias entre categorías ideológicas',
        style='List Number'
    )

    doc.add_heading('3.3. Resultados: Ítems Significativos', 2)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = ['Elección', 'CT Significativos', '% del Total']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ['Generales', '14 / 40', '35%'],
        ['Ballotage', '5 / 40', '12.5%']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    p = doc.add_paragraph()
    p.add_run('Observación importante: ').bold = True
    p.add_run(
        'Generales mostró casi 3 veces más ítems CT significativos que Ballotage, '
        'sugiriendo mayor diversidad de procesamiento ideológico en la primera vuelta electoral.'
    )

    doc.add_heading('3.4. Variables CT_Filtrado Creadas', 2)

    p = doc.add_paragraph()
    p.add_run('Se crearon 4 variables por elección, representando el cruce Tipo de Ítem × Ideología Candidato:')

    doc.add_paragraph('Cambio_Tiempo_Filt_Pro_Izq: Ítems Progresistas × Candidatos Izquierda', style='List Bullet')
    doc.add_paragraph('Cambio_Tiempo_Filt_Pro_Der: Ítems Progresistas × Candidatos Derecha', style='List Bullet')
    doc.add_paragraph('Cambio_Tiempo_Filt_Con_Izq: Ítems Conservadores × Candidatos Izquierda', style='List Bullet')
    doc.add_paragraph('Cambio_Tiempo_Filt_Con_Der: Ítems Conservadores × Candidatos Derecha', style='List Bullet')

    doc.add_heading('3.5. Análisis Post-Hoc', 2)

    p = doc.add_paragraph()
    p.add_run(
        'Se realizaron comparaciones pareadas (test de Wilcoxon) entre las variables filtradas '
        'para identificar patrones de procesamiento diferencial.'
    )

    p = doc.add_paragraph()
    p.add_run('Hallazgos principales:')

    doc.add_paragraph(
        'Las diferencias entre Pro_Izq y Pro_Der fueron significativas en Generales pero no en Ballotage',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Los tiempos para ítems conservadores mostraron mayor asimetría ideológica',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Las variables filtradas mostraron efectos de mayor magnitud que las variables sumadas completas',
        style='List Bullet'
    )

    doc.add_heading('3.6. Archivos Generados', 2)

    doc.add_paragraph('/Data/Procesados/Generales.xlsx (con variables CT_Filt)', style='List Bullet')
    doc.add_paragraph('/Data/Procesados/Ballotage.xlsx (con variables CT_Filt)', style='List Bullet')
    doc.add_paragraph('/Data/Resultados_Cleveland/Resumen_Resultados_CT_Filtrado.xlsx', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Notebooks fuente: ').bold = True
    p.add_run('37, 38, 39, 51 (serie completa de CT Filtrados)')

    doc.add_page_break()

    print("Continuando con Sección 4...")

    # ===========================================================================
    # SECCIÓN 4: DIFERENCIA DE DIFERENCIAS
    # ===========================================================================

    doc.add_heading('4. Diferencia de Diferencias (DifDif)', 1)

    doc.add_heading('4.1. Marco Conceptual', 2)

    p = doc.add_paragraph()
    p.add_run('Objetivo: ').bold = True
    p.add_run(
        'Cuantificar cómo cambia la asimetría ideológica (diferencia Izquierda-Derecha) '
        'entre la primera y segunda vuelta electoral.'
    )

    p = doc.add_paragraph()
    p.add_run('Fórmula:')

    doc.add_paragraph('Dif_Gen = Izquierda - Derecha (en Generales)', style='List Bullet')
    doc.add_paragraph('Dif_Bal = Izquierda - Derecha (en Ballotage)', style='List Bullet')
    doc.add_paragraph('DifDif = Dif_Bal - Dif_Gen', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Interpretación:')

    doc.add_paragraph('DifDif > 0: La asimetría Izq-Der aumentó en Ballotage', style='List Bullet')
    doc.add_paragraph('DifDif < 0: La asimetría Izq-Der disminuyó en Ballotage', style='List Bullet')
    doc.add_paragraph('DifDif = 0: No hubo cambio en la asimetría', style='List Bullet')

    doc.add_heading('4.2. Variables Calculadas', 2)

    p = doc.add_paragraph()
    p.add_run('Para cada uno de los 20 ítems, se calcularon 3 variables:')

    doc.add_paragraph('Dif_Gen_[CO/CT]_Item_X: Diferencia Izq-Der en Generales', style='List Bullet')
    doc.add_paragraph('Dif_Bal_[CO/CT]_Item_X: Diferencia Izq-Der en Ballotage', style='List Bullet')
    doc.add_paragraph('DifDif_[CO/CT]_Item_X: Cambio en la asimetría entre elecciones', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Total: 120 variables (60 para CO, 60 para CT)')

    doc.add_heading('4.3. Características de los Datos', 2)

    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = ['Tipo', 'n Promedio', 'Media DifDif', 'DE DifDif']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ['DifDif_CO', '~155', '~0.05', '~1.5'],
        ['DifDif_CT', '~155', '~1.2', '~11.0']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    p = doc.add_paragraph()
    p.add_run('Nota: ').italic = True
    p.add_run(
        'Los valores son promedios aproximados basados en las primeras 5 variables de cada tipo. '
        'Los valores faltantes (~2620 por variable) se deben a que solo los participantes que '
        'respondieron en ambas elecciones tienen valores DifDif calculables.'
    )

    doc.add_heading('4.4. Gráficos de Cleveland Generados', 2)

    p = doc.add_paragraph()
    p.add_run(
        'Se generaron gráficos de Cleveland (dot plots) para visualizar las diferencias entre elecciones:'
    )

    doc.add_paragraph(
        'Notebook 60: CT Diferencias para todos los ítems combinados',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Notebook 61: CT Diferencias separadas por tipo de ítem (Progresistas vs Conservadores)',
        style='List Bullet'
    )

    p = doc.add_paragraph()
    p.add_run('Características de los gráficos:')

    doc.add_paragraph('Punto azul: Diferencia Izq-Der en Generales', style='List Bullet')
    doc.add_paragraph('Punto rojo: Diferencia Izq-Der en Ballotage', style='List Bullet')
    doc.add_paragraph('Línea conectora: Muestra el cambio (DifDif)', style='List Bullet')
    doc.add_paragraph('Asteriscos: Indican significancia estadística del cambio', style='List Bullet')

    doc.add_heading('4.5. Análisis de Significancia', 2)

    p = doc.add_paragraph()
    p.add_run('Metodología: ').bold = True
    p.add_run(
        'Test t pareado para determinar si el cambio en asimetría (DifDif) es '
        'estadísticamente diferente de cero para cada ítem.'
    )

    p = doc.add_paragraph()
    p.add_run('Criterios de significancia:')

    doc.add_paragraph('* p < 0.05', style='List Bullet')
    doc.add_paragraph('** p < 0.01', style='List Bullet')
    doc.add_paragraph('*** p < 0.001', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Interpretación: ').bold = True
    p.add_run(
        'Los ítems con cambios significativos representan temas donde la polarización ideológica '
        'aumentó o disminuyó sustancialmente entre primera y segunda vuelta.'
    )

    doc.add_heading('4.6. Archivos Generados', 2)

    doc.add_paragraph('/Data/Procesados/df_Elecciones.xlsx (con variables DifDif)', style='List Bullet')
    doc.add_paragraph('/Código/Graficos_Cleveland/Cleveland_CT_Todos_Items.png', style='List Bullet')
    doc.add_paragraph('/Código/Graficos_Cleveland/Cleveland_CT_Por_Tipo_Panel_A.png', style='List Bullet')
    doc.add_paragraph('/Código/Graficos_Cleveland/Cleveland_CT_Por_Tipo_Panel_B.png', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Notebooks fuente: ').bold = True
    p.add_run('48 (cálculos), 60 y 61 (visualizaciones)')

    doc.add_page_break()

    print("Continuando con Sección 5...")

    # ===========================================================================
    # SECCIÓN 5: MODELOS SEM
    # ===========================================================================

    doc.add_heading('5. Modelos de Ecuaciones Estructurales (SEM)', 1)

    doc.add_heading('5.1. Marco Conceptual', 2)

    p = doc.add_paragraph()
    p.add_run('Objetivo general: ').bold = True
    p.add_run(
        'Determinar en qué medida los índices ideológicos (Progresismo y Conservadurismo) '
        'predicen los cambios de opinión y tiempo hacia candidatos de izquierda y derecha.'
    )

    p = doc.add_paragraph()
    p.add_run('Pregunta de investigación: ').italic = True
    p.add_run(
        '¿La ideología del participante predice su patrón de cambios de opinión y tiempos de respuesta '
        'cuando evalúa a candidatos de diferentes orientaciones políticas?'
    )

    doc.add_heading('5.2. Especificación de los Modelos', 2)

    p = doc.add_paragraph()
    p.add_run('Estructura general:')

    doc.add_paragraph(
        'Variables predictoras (exógenas): Indice_Progresismo, Indice_Conservadurismo',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Variables outcome (endógenas): Variables de CO y CT',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Método de estimación: Mínimos Cuadrados Generalizados (GLS)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Software: Python con librería semopy',
        style='List Bullet'
    )

    doc.add_heading('5.3. Grupos de Modelos Analizados', 2)

    p = doc.add_paragraph()
    p.add_run('Se ejecutaron 4 grupos de modelos, cada uno en Generales y Ballotage:')

    # Tabla de modelos
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = ['Grupo', 'Variables Outcome', 'N Modelos', 'Notebook']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ['Variables Sumadas', 'CO_Sum y CT_Sum (4 vars c/u)', '8', '55'],
        ['Variables Filtradas', 'CO_Filt y CT_Filt (4 vars c/u)', '8', '56'],
        ['Congruencia', 'CO/CT Congruente vs Incongruente', '4', '57'],
        ['Por Tipo Ítem', 'Total Progresistas vs Conservadores', '4', '58']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    p = doc.add_paragraph()
    p.add_run('Total: 24 modelos × 2 elecciones = 48 modelos SEM')

    doc.add_heading('5.4. Métricas de Ajuste Utilizadas', 2)

    doc.add_paragraph('R²: Proporción de varianza explicada (0-1, mayor es mejor)', style='List Bullet')
    doc.add_paragraph('AIC (Akaike Information Criterion): Menor es mejor', style='List Bullet')
    doc.add_paragraph('BIC (Bayesian Information Criterion): Menor es mejor', style='List Bullet')
    doc.add_paragraph('CFI (Comparative Fit Index): > 0.95 indica buen ajuste', style='List Bullet')
    doc.add_paragraph('TLI (Tucker-Lewis Index): > 0.95 indica buen ajuste', style='List Bullet')
    doc.add_paragraph('RMSEA (Root Mean Square Error): < 0.06 indica buen ajuste', style='List Bullet')
    doc.add_paragraph('SRMR (Standardized Root Mean Square Residual): < 0.08 indica buen ajuste', style='List Bullet')

    doc.add_heading('5.5. Resultados Principales', 2)

    # Top 5 modelos
    doc.add_heading('5.5.1. Modelos con Mejor Ajuste (Top 5 por R²)', 3)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = ['Ranking', 'Elección', 'Outcome', 'R²']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ['1', 'Ballotage', 'Cambio_Op_Sum_Pro_Der', '0.0449'],
        ['2', 'Ballotage', 'Cambio_Op_Sum_Con_Izq', '0.0437'],
        ['3', 'Ballotage', 'Cambio_Op_Sum_Pro_Izq', '0.0419'],
        ['4', 'Generales', 'Cambio_Op_Sum_Con_Izq', '0.0358'],
        ['5', 'Generales', 'Cambio_Op_Sum_Pro_Izq', '0.0352']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    p = doc.add_paragraph()
    p.add_run('Observaciones:')

    doc.add_paragraph(
        'Los modelos para Cambio de Opinión (CO) sistemáticamente superan a los de Cambio de Tiempo (CT)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Ballotage muestra R² ligeramente superiores a Generales',
        style='List Bullet'
    )
    doc.add_paragraph(
        'El mejor modelo explica 4.49% de la varianza, un valor modesto pero estadísticamente significativo',
        style='List Bullet'
    )

    doc.add_heading('5.5.2. Poder Predictivo por Tipo de Variable', 3)

    table = doc.add_table(rows=2, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = ['Tipo de Variable', 'R² Promedio (Generales)', 'R² Promedio (Ballotage)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ['Variables Sumadas', '0.0181 ± 0.016', '0.0209 ± 0.020']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    p = doc.add_paragraph()
    p.add_run('Nota: ').italic = True
    p.add_run(
        'Solo se pudo analizar el grupo "Variables Sumadas" debido a que los archivos '
        'de resultados de otros grupos no estaban disponibles en el momento del análisis.'
    )

    doc.add_heading('5.5.3. Predictores Más Importantes', 3)

    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = ['Elección', 'Predictor', 'β Promedio', '% Significativos']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ['Generales', 'Indice_Conservadurismo', '-1.019', '75%'],
        ['Ballotage', 'Indice_Progresismo', '0.377', '50%']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    p = doc.add_paragraph()
    p.add_run('Interpretación:')

    doc.add_paragraph(
        'En Generales, el Índice de Conservadurismo es el predictor dominante (coeficiente β mayor y 75% de modelos con efectos significativos)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'En Ballotage, ambos índices muestran efectos similares en términos de frecuencia de significancia (50%)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'El signo negativo del Conservadurismo sugiere que mayor conservadurismo se asocia con menores cambios en ciertas direcciones',
        style='List Bullet'
    )

    doc.add_heading('5.6. Limitaciones del Análisis SEM', 2)

    doc.add_paragraph(
        'Baja proporción de varianza explicada: El R² promedio de ~2% indica que otros factores no incluidos en el modelo (contexto electoral, características demográficas, exposición a medios) probablemente juegan roles importantes',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Datos faltantes: No todos los participantes completaron ambas rondas electorales, reduciendo el tamaño muestral para análisis longitudinales',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Multicolinealidad potencial: Los índices de Progresismo y Conservadurismo pueden estar correlacionados negativamente, aunque esto es teóricamente esperado',
        style='List Bullet'
    )

    doc.add_heading('5.7. Conclusiones del Análisis SEM', 2)

    p = doc.add_paragraph()
    p.add_run('Pregunta: ').bold = True
    p.add_run('¿La ideología predice cambios de opinión y tiempo?')

    p = doc.add_paragraph()
    p.add_run('Respuesta: ').bold = True
    p.add_run(
        'Sí, pero con efectos pequeños. Los modelos SEM demuestran que la ideología política '
        '(medida mediante índices de Progresismo y Conservadurismo) predice significativamente '
        'los cambios de opinión hacia candidatos, aunque la magnitud del efecto es modesta (R² < 5%).'
    )

    p = doc.add_paragraph()
    p.add_run('Implicaciones:')

    doc.add_paragraph(
        'Los cambios de opinión son fenómenos multifactoriales que no dependen únicamente de la ideología previa',
        style='List Bullet'
    )
    doc.add_paragraph(
        'El contexto electoral (Generales vs Ballotage) modera la relación ideología-cambio',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Los cambios de tiempo de respuesta parecen menos influenciados por ideología que los cambios de opinión',
        style='List Bullet'
    )

    doc.add_heading('5.8. Archivos Generados', 2)

    doc.add_paragraph('/Data/Resultados_SEM/SEM_Variables_Sumadas_Generales.xlsx', style='List Bullet')
    doc.add_paragraph('/Data/Resultados_SEM/SEM_Variables_Sumadas_Ballotage.xlsx', style='List Bullet')
    doc.add_paragraph('/Data/Resultados_SEM/RESUMEN_CONSOLIDADO_SEM.xlsx', style='List Bullet')
    doc.add_paragraph('/Data/Resultados_SEM/Heatmap_Correlaciones_Generales.png', style='List Bullet')
    doc.add_paragraph('/Data/Resultados_SEM/Heatmap_Correlaciones_Ballotage.png', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Notebooks fuente: ').bold = True
    p.add_run('54 (Correlaciones), 55-58 (Modelos), 59 (Resumen y Comparación)')

    print("Documento completado...")

    # Guardar documento
    return doc

if __name__ == "__main__":
    doc = crear_documento()
    output_path = '/home/user/Tesis/Resumen_Analisis_Estadisticos.docx'
    doc.save(output_path)

    print(f"\n✅ Documento generado exitosamente: {output_path}")
    print(f"Tamaño aproximado: 15-18 páginas")
